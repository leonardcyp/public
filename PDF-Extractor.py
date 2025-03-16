import os
import customtkinter as ctk
import fitz  # PyMuPDF for PDF processing
from PIL import Image, ImageTk
import pytesseract
from tkinter import filedialog, messagebox
from tkinter import Canvas, NW
from docx import Document  # Import the python-docx library

# Set dark appearance for a modern look
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

# -----------------------------
# Main Application Class
# -----------------------------
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("DeepSeek Text Extraction Application")
        self.geometry("1200x800")
        self.resizable(True, True)

        container = ctk.CTkFrame(self, fg_color="#000000")
        container.pack(fill="both", expand=True)

        self.frames = {}
        # Only the HomePage is used now.
        for F in (HomePage,):
            frame = F(container, self)
            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(HomePage)

    def show_frame(self, page_class):
        frame = self.frames[page_class]
        frame.tkraise()

# -----------------------------
# Home Page
# -----------------------------
class HomePage(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="#000000")
        self.controller = controller

        # File & extraction variables
        self.file_path = None
        self.extracted_text = ""
        self.preview_image = None   # For preview display
        self.original_img = None    # Full-resolution PIL Image
        self.preview_img_width = None  # Width of the scaled preview image
        self.preview_img_height = None # Height of the scaled preview image

        # Variables for interactive extraction area selection
        self.start_x = None
        self.start_y = None
        self.end_x = None
        self.end_y = None
        self.rect = None

        # Scrollable frame for content
        self.scroll_frame = ctk.CTkScrollableFrame(self, fg_color="#000000", width=1150, height=700)
        self.scroll_frame.pack(fill="both", expand=True, padx=10, pady=1)

        # Title
        title = ctk.CTkLabel(self.scroll_frame, text="Home Page", font=("Arial", 24), anchor="nw")
        title.pack(pady=(10, 20), anchor="nw")

        # File row: path & upload button
        file_frame = ctk.CTkFrame(self.scroll_frame, fg_color="#000000")
        file_frame.pack(pady=5, anchor="nw")

        file_label = ctk.CTkLabel(file_frame, text="Uploaded File:", anchor="nw")
        file_label.pack(side="left", padx=(0, 10))

        self.file_path_entry = ctk.CTkEntry(file_frame, width=600)
        self.file_path_entry.configure(state="readonly")
        self.file_path_entry.pack(side="left", padx=(0, 10))

        upload_btn = ctk.CTkButton(file_frame, text="Upload PDF/Image File", command=self.upload_file, width=200)
        upload_btn.pack(side="left")

        # Preview area (enlarged)
        left_frame = ctk.CTkFrame(self.scroll_frame, fg_color="#000000")
        left_frame.pack(pady=20, anchor="nw", padx=10)

        preview_label = ctk.CTkLabel(left_frame, text="Original File Preview (Select Extraction Area):", anchor="nw")
        preview_label.pack(pady=(0, 5), anchor="nw")

        # Enlarged preview canvas
        self.preview_canvas = Canvas(left_frame, width=600, height=600, bg="#000000", highlightthickness=0)
        self.preview_canvas.pack(pady=(0, 5), anchor="nw", padx=(300, 0))
        # Bind mouse events for selection on the preview canvas
        self.preview_canvas.bind("<ButtonPress-1>", self.on_button_press)
        self.preview_canvas.bind("<B1-Motion>", self.on_move_press)
        self.preview_canvas.bind("<ButtonRelease-1>", self.on_button_release)

        # Extracted text display
        extracted_label = ctk.CTkLabel(self.scroll_frame, text="Extracted Text:", anchor="nw")
        extracted_label.pack(pady=(10, 2), anchor="nw")

        self.text_display = ctk.CTkTextbox(self.scroll_frame, width=800, height=200, corner_radius=8)
        self.text_display.pack(pady=(0, 10), anchor="nw")

        # Button for text extraction
        btn_frame = ctk.CTkFrame(self.scroll_frame, fg_color="#000000")
        btn_frame.pack(pady=5, anchor="nw")

        extract_btn = ctk.CTkButton(btn_frame, text="Extract Text", command=self.extract_text, width=200)
        extract_btn.pack(side="left", padx=(0, 10))

        # Button for text extraction and export to Word
        export_btn = ctk.CTkButton(btn_frame, text="Extract & Export to Word", command=self.extract_and_export_to_word, width=200)
        export_btn.pack(side="left", padx=(0, 10))

    # ----------------- File Upload -----------------
    def upload_file(self):
        filetypes = [("PDF files", "*.pdf"), ("Image files", "*.jpg *.jpeg *.png")]
        self.file_path = filedialog.askopenfilename(title="Select a file", filetypes=filetypes)
        if self.file_path:
            self.file_path_entry.configure(state="normal")
            self.file_path_entry.delete(0, "end")
            self.file_path_entry.insert(0, self.file_path)
            self.file_path_entry.configure(state="readonly")
            self.display_original_file()
            # Reset any previous selection
            self.start_x = self.start_y = self.end_x = self.end_y = None
            if self.rect:
                self.preview_canvas.delete(self.rect)
                self.rect = None

    def display_original_file(self):
        """Display an enlarged preview of the uploaded file."""
        self.preview_canvas.delete("all")
        ext = os.path.splitext(self.file_path)[1].lower()
        try:
            if ext == ".pdf":
                doc = fitz.open(self.file_path)
                page = doc.load_page(0)
                pix = page.get_pixmap()
                mode = "RGB" if pix.n < 4 else "RGBA"
                full_img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                doc.close()
            elif ext in [".jpg", ".jpeg", ".png"]:
                full_img = Image.open(self.file_path)
            else:
                messagebox.showerror("Unsupported File", "The file type is not supported for preview.")
                return

            self.original_img = full_img

            # Create an enlarged preview copy
            preview_img = full_img.copy()
            preview_img.thumbnail((600, 600))
            self.preview_img_width, self.preview_img_height = preview_img.size
            self.preview_image = ImageTk.PhotoImage(preview_img)
            self.preview_canvas.create_image(0, 0, anchor=NW, image=self.preview_image)

        except Exception as e:
            messagebox.showerror("Preview Error", f"Error displaying file preview: {e}")

    # ----------------- Mouse Event Handlers for Extraction Area -----------------
    def on_button_press(self, event):
        # Record the starting coordinates for selection
        self.start_x = event.x
        self.start_y = event.y
        # Delete previous rectangle if it exists
        if self.rect:
            self.preview_canvas.delete(self.rect)
        self.rect = self.preview_canvas.create_rectangle(self.start_x, self.start_y, self.start_x, self.start_y, outline='red', width=2)

    def on_move_press(self, event):
        # Update the rectangle as the mouse is dragged
        curX, curY = event.x, event.y
        self.preview_canvas.coords(self.rect, self.start_x, self.start_y, curX, curY)

    def on_button_release(self, event):
        # Finalize the selection area
        self.end_x = event.x
        self.end_y = event.y

    # ----------------- Text Extraction -----------------
    def extract_text(self):
        if not self.file_path:
            messagebox.showwarning("No File", "Please upload a file first!")
            return
        if not self.original_img:
            messagebox.showerror("Image Error", "Original image not loaded properly.")
            return

        try:
            if self.start_x is not None and self.end_x is not None:
                # Get the selected coordinates from the preview canvas
                x1, y1 = min(self.start_x, self.end_x), min(self.start_y, self.end_y)
                x2, y2 = max(self.start_x, self.end_x), max(self.start_y, self.end_y)

                # Compute the scaling factors between the original image and the preview
                scale_x = self.original_img.width / self.preview_img_width
                scale_y = self.original_img.height / self.preview_img_height

                # Map the preview coordinates to the original image coordinates
                orig_x1 = int(x1 * scale_x)
                orig_y1 = int(y1 * scale_y)
                orig_x2 = int(x2 * scale_x)
                orig_y2 = int(y2 * scale_y)

                roi = self.original_img.crop((orig_x1, orig_y1, orig_x2, orig_y2))
                self.extracted_text = pytesseract.image_to_string(roi, config='--oem 3 --psm 6')
            else:
                ext = os.path.splitext(self.file_path)[1].lower()
                if ext == ".pdf":
                    doc = fitz.open(self.file_path)
                    text = ""
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text += page.get_text()
                    self.extracted_text = text
                else:
                    self.extracted_text = pytesseract.image_to_string(self.original_img, config='--oem 3 --psm 6')

            self.text_display.delete("1.0", "end")
            self.text_display.insert("end", self.extracted_text)

        except Exception as e:
            messagebox.showerror("Extraction Error", f"Error extracting text: {e}")

    # ----------------- Extract and Export to Word -----------------
    def extract_and_export_to_word(self):
        self.extract_text()
        if self.extracted_text:
            try:
                doc = Document()
                doc.add_heading('Extracted Text', level=1)
                doc.add_paragraph(self.extracted_text)
                save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
                if save_path:
                    doc.save(save_path)
                    messagebox.showinfo("Success", "Text successfully exported to Word file.")
            except Exception as e:
                messagebox.showerror("Export Error", f"Error exporting text to Word: {e}")

# -----------------------------
# Main Execution
# -----------------------------
if __name__ == "__main__":
    app = App()
    app.mainloop()
